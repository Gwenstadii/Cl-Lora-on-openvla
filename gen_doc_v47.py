from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Calibri'

# ============ Title ============
doc.add_heading('CL-LoRA Prototype Replay V47 实验报告', level=0)
doc.add_paragraph('')

# ============ 1. V47 配置 ============
doc.add_heading('1. V47 配置', level=1)

doc.add_heading('1.1 架构', level=2)
doc.add_paragraph(
    'V47 基于 PI 式 Action Expert 设计，CL-LoRA 仅注入末尾 12 层 (L20-31)，'
    '前 20 层为 bare frozen backbone。'
)
p = doc.add_paragraph()
p.add_run('注入结构：').bold = True
doc.add_paragraph('L0-19:  bare backbone (frozen, 无 LoRA)', style='List Bullet')
doc.add_paragraph('L20-23: shared LoRA (4 层, Stage 1 后 A+B 冻结)', style='List Bullet')
doc.add_paragraph('L24-31: specific LoRA (8 层, B+block 存 bank)', style='List Bullet')
doc.add_paragraph('Action Head: 注入 CL-LoRA (全 specific), B+block 存 bank', style='List Bullet')

doc.add_heading('1.2 参数配置', level=2)
table = doc.add_table(rows=11, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
config = [
    ('Model', 'OpenVLA-7B (LLaMA-2-7B + DinoSigLIP)'),
    ('LoRA rank', '16'),
    ('first_lora_layer', '20'),
    ('shared_depth', '4 (L20-23 shared, L24-31 specific)'),
    ('freeze_specific_a', 'False (特定 A 跨任务漂移)'),
    ('Action Head LoRA', '全 specific, rank=16'),
    ('Bank 内容', 'Specific B + block_scale (backbone + action_head)'),
    ('学习率', '5e-4, warmup 200 steps'),
    ('Stage 1 步数', '12000'),
    ('Stage 2-4 步数', '16000'),
]
for i, (k, v) in enumerate(config):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    if i == 0:
        for c in range(2):
            for p in table.rows[0].cells[c].paragraphs:
                for r in p.runs:
                    r.bold = True
doc.add_paragraph('')

doc.add_heading('1.3 任务序列', level=2)
table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Stage', '任务', '数据集']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
for i, (s, t, d) in enumerate([('1','A','libero_spatial'),('2','B','libero_spatial'),
                                  ('3','C','libero_object'),('4','D','libero_goal')]):
    table.rows[i+1].cells[0].text = s
    table.rows[i+1].cells[1].text = t
    table.rows[i+1].cells[2].text = d

# ============ 2. 结果 ============
doc.add_heading('2. 实验结果', level=1)

doc.add_heading('2.1 V47 无回放', level=2)
table = doc.add_table(rows=5, cols=5, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['', 'Stage 1', 'Stage 2', 'Stage 3', 'Stage 4']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
for i, h in enumerate(['步数', '12000', '16000', '16000', '16000']):
    table.rows[1].cells[i].text = h
data = [
    ('Task A', '0.96', '0.52', '0.08', '0.00'),
    ('Task B', '—', '0.84', '0.42', '0.30'),
    ('Task C', '—', '—', '1.00', '0.60'),
    ('Task D', '—', '—', '—', '1.00'),
]
for i, (label, s1, s2, s3, s4) in enumerate(data):
    table.rows[i+1].cells[0].text = label
    for j, v in enumerate([s1, s2, s3, s4]):
        table.rows[i+1].cells[j+1].text = v
    for p in table.rows[i+1].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True

doc.add_paragraph('')
doc.add_paragraph('遗忘模式：Stage 2 (A=0.52) → Stage 3 (A=0.08, B=0.42) → Stage 4 (A=0, B=0.30, C=0.60)。'
                   '跨域时遗忘加速，A 经历 3 步漂移后全崩，C 经历 1 步漂移降至 0.60。')

doc.add_heading('2.2 V47 Replay', level=2)
table = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['', 'Stage 2', 'Stage 3', 'Stage 4']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
for i, h in enumerate(['步数', '16000', '16000', '16000']):
    table.rows[1].cells[i].text = h
data = [
    ('Task A', '1.00', '0.22', '0.00'),
    ('Task B', '0.90', '0.90', '0.98'),
    ('Task C', '—', '1.00', '0.92'),
    ('Task D', '—', '—', '1.00'),
]
for i, (label, s2, s3, s4) in enumerate(data):
    table.rows[i+1].cells[0].text = label
    for j, v in enumerate([s2, s3, s4]):
        table.rows[i+1].cells[j+1].text = v
    for p in table.rows[i+1].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True

doc.add_paragraph('')
doc.add_paragraph('Replay 效果：B/C 几乎完美恢复 (≥0.90)，D 正常学习 (1.00)。A 在 Stage 3 仅 0.22，Stage 4 降至 0。')

doc.add_heading('2.3 无回放 vs Replay 对比 (Stage 4)', level=2)
table = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Task', '无回放', 'Replay', 'Gap']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
for i, (t, nr, r, g) in enumerate([
    ('A', '0.00', '0.00', '—'),
    ('B', '0.30', '0.98', '+0.68'),
    ('C', '0.60', '0.92', '+0.32'),
    ('D', '1.00', '1.00', '—'),
]):
    table.rows[i+1].cells[0].text = t
    table.rows[i+1].cells[1].text = nr
    table.rows[i+1].cells[2].text = r
    table.rows[i+1].cells[3].text = g

# ============ 3. 问题分析 ============
doc.add_heading('3. 存在的问题与原因', level=1)

doc.add_heading('3.1 Task A Retention 极低 (0.00-0.22)', level=2)
doc.add_paragraph(
    'Task A (libero_spatial) 经历了 3 个后续任务的连续训练 (B→C→D)，'
    '特定 A 累积漂移 = 12000(A) + 16000(B) + 16000(C) + 16000(D) = 60000 步。'
    '8 层特定 A (L24-31) 已面目全非，bank 恢复的 specific B 与当前特定 A 完全不兼容。'
)
doc.add_paragraph(
    'Replay 拉力 = 0.5 weight × 1/15 frequency = 3% 有效梯度，'
    '无法对抗 60000 步累积漂移。A 是所有任务中最老的，replay buffer 信号最稀疏。',
    style='List Bullet'
)
doc.add_paragraph(
    '对比 V38 (freeze_specific_a=True, 特定 A 永冻)：Replay 可将 A 从 0.44 拉到 0.96。'
    'freeze_specific_a 是 replay 能否拉回 A 的关键前提。',
    style='List Bullet'
)

doc.add_heading('3.2 B/C Replay 效果显著', level=2)
doc.add_paragraph(
    'Task B 从无回放的 0.30 提升到 replay 的 0.98 (+0.68)。'
    'Task C 从无回放的 0.60 提升到 replay 的 0.92 (+0.32)。'
    '这说明 prototype replay 对最近 2 个任务的保护极为有效——'
    '漂移步数少 (B=32000 步, C=16000 步)，replay 梯度足够抗衡。'
)

doc.add_heading('3.3 freeze_specific_a 的矛盾', level=2)
p = doc.add_paragraph()
p.add_run('freeze_specific_a=True (V38): ').bold = True
p.add_run('特定 A 永冻 → 无回放 baseline 偏高 (A=0.44) → replay gap 虽大，论文可接受')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('freeze_specific_a=False (V47): ').bold = True
p.add_run('特定 A 漂移 → 无回放 baseline 极低 (A=0) → 但 replay 也拉不回最老任务 → 只有近 2 个任务有效')

doc.add_heading('3.4 当前最优方案', level=2)
doc.add_paragraph(
    'V38 (first_lora_layer=24, shared=4, specific=4, freeze_specific_a=True, rank=16)：\n'
    '  无回放: 0.96 | 0.42-0.80 | 0.66-0.90-0.98 | 0.44-0.66-0.94-1.00\n'
    '  Replay:   —  | 0.96-0.94  | 0.96-0.94-0.88 | 0.96-0.94-0.88-1.00'
)
doc.add_paragraph(
    'V48 (first_lora_layer=20, shared=4, specific=8, freeze_specific_a=True, rank=16) 训练中：\n'
    '  预期无回放 baseline 介于 V38-V47 之间，replay 能拉回全部任务。'
)

# ============ 4. 总结 ============
doc.add_heading('4. 总结', level=1)
bullets = [
    'PI 式 Action Expert 架构 (first_lora_layer) 成功将 CL-LoRA 从 32 层 backbone 收缩到末尾 8-12 层',
    'freeze_specific_a 是 replay 有效性的关键：True → 旧任务 subspace 稳定 → replay 拉力足够恢复',
    'Prototype Replay v2 对最近 1-2 个任务极其有效 (≥0.90)，最老任务 (A) 需要更强的 replay 策略',
    'V38 为目前最优结果：单任务 0.96，无回放跨域最低 0.42，Replay ≥0.88',
    'V48 正在训练中，预期进一步拉大无回放 vs replay 的 gap',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.save(r'C:\Users\pengs\Documents\GitHub\Cl-Lora-on-openvla\1st.docx')
print('Done: 1st.docx')
